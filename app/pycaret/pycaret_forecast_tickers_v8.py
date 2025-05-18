#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para entrenamiento automático de modelos de series temporales por ticker y variable.
Genera resultados, gráficos y reporte DOCX en un directorio único por corrida.
"""
import os
import sys
import logging
import datetime
from pathlib import Path
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from sqlalchemy import create_engine
from pycaret.time_series import *

# ========== PARAMETRIZACIÓN ==========
PARAMS = {
    'predict_precio_cierre': True,
    'predict_precio_max':    False,
    'predict_precio_min':    False,
    'predict_volumen_operado': False,
}
TARGET_VARS = [var for var, flag in {
    'precio_cierre':        PARAMS['predict_precio_cierre'],
    'precio_max':           PARAMS['predict_precio_max'],
    'precio_min':           PARAMS['predict_precio_min'],
    'volumen_operado':      PARAMS['predict_volumen_operado'],
}.items() if flag]
FH = 22  # horizonte de predicción (días)

#Modelos a probar (por defecto Prophet y RandomForestRegressor)
MODELOS_A_PROBAR = ['prophet', 'rf_cds_dt']  # Puedes poner None para probar todos

# Fecha de inicio para considerar los registros (inclusive)
FECHA_INICIO_DATOS = '2024-02-01'

# ========== LOGGING Y DIRECTORIOS ==========
RUN_TS = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
OUT_DIR = Path(f'outputs_{RUN_TS}')
OUT_DIR.mkdir(parents=True, exist_ok=True)

logging.basicConfig(
    filename=OUT_DIR/'run.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)
logger.info(f'Variables a modelar: {TARGET_VARS}')
print(f'Resultados se guardarán en: {OUT_DIR}')

# ========== CONEXIÓN BASE DE DATOS ==========
USER = os.getenv("DB_USER", "usuario")
PASSWORD = os.getenv("DB_PASSWORD", "contraseña")
HOST = os.getenv("DB_HOST", "db")
DB = os.getenv("DB_NAME", "patagonia_db")
DATABASE_URL = f"postgresql://{USER}:{PASSWORD}@{HOST}/{DB}"
engine = create_engine(DATABASE_URL)
print(f'Conectado a: {DATABASE_URL}')

# ========== CARGA DE DATOS ==========
query = '''
SELECT c.id_empresa,
       e.ticker,
       c.fecha,
       c.precio_apertura,
       c.precio_cierre,
       c.precio_max,
       c.precio_min,
       c.volumen_operado,
       c.variacion_porcentaje
FROM cotizacion_x_empresa c
JOIN empresas e ON c.id_empresa = e.id_empresa
ORDER BY e.ticker, c.fecha;
'''
df = pd.read_sql(query, engine)
df['fecha'] = pd.to_datetime(df['fecha'])
for col in ['precio_apertura','precio_cierre','precio_max','precio_min','volumen_operado','variacion_porcentaje']:
    df[col] = pd.to_numeric(df[col], errors='coerce')
# Filtrar por fecha de inicio
if FECHA_INICIO_DATOS:
    df = df[df['fecha'] >= FECHA_INICIO_DATOS]
print('Dataset shape:', df.shape)

# ========== FUNCIONES AUXILIARES ==========
def plot_forecast(original: pd.Series, forecast_df: pd.DataFrame, title:str, zoom_window:int=30, model_name:str=None):
    preds = forecast_df['y_pred']
    lower = forecast_df.get('y_pred_lower')
    upper = forecast_df.get('y_pred_upper')
    def _draw(ax, hist_idx, subtitle):
        ax.plot(hist_idx, original.loc[hist_idx].values, label='Histórico')
        ax.plot(preds.index, preds.values, linestyle='--', marker='o', label='Pronosticado')
        if lower is not None and upper is not None and lower.notna().all():
            ax.fill_between(preds.index, lower.values, upper.values, alpha=0.25, label='IC 95 %')
        ax.set_title(subtitle)
        ax.grid(True)
        ax.legend()
    # gráfico completo
    model_str = f" [{model_name}]" if model_name else ""
    fig, ax = plt.subplots(figsize=(11,4))
    _draw(ax, original.index, f'{title}{model_str} – completo')
    fig.savefig(OUT_DIR/f'{title.replace(" / ", "_")}_full.png')
    plt.close(fig)
    # zoom
    zoom_start = original.index[-zoom_window] if len(original) >= zoom_window else original.index[0]
    fig, ax = plt.subplots(figsize=(11,4))
    _draw(ax, original.loc[zoom_start:].index, f'{title}{model_str} – últimos {zoom_window} d + forecast')
    fig.savefig(OUT_DIR/f'{title.replace(" / ", "_")}_zoom.png')
    plt.close(fig)

def train_forecast(series: pd.Series, fh:int = 22, sort_metric:str = 'MASE', modelos_a_probar=None):
    series = series.sort_index()
    freq = pd.infer_freq(series.index) or 'B'
    series = series.asfreq(freq).ffill().bfill()
    exp = setup(
        data=series,
        fh=fh,
        fold=3,
        session_id=123,
        verbose=False
    )
    if modelos_a_probar:
        best = compare_models(include=modelos_a_probar, sort=sort_metric, verbose=False)
    else:
        best = compare_models(sort=sort_metric, verbose=False)
    preds_df = predict_model(best, fh=fh, coverage=[0.025, 0.975])
    lower_cols = [c for c in preds_df.columns if 'lower' in c.lower()]
    upper_cols = [c for c in preds_df.columns if 'upper' in c.lower()]
    if lower_cols and upper_cols:
        preds_df = preds_df.rename(columns={lower_cols[0]:'y_pred_lower', upper_cols[0]:'y_pred_upper'})
    else:
        preds_df['y_pred_lower'] = pd.NA
        preds_df['y_pred_upper'] = pd.NA
    return best, preds_df

# ========== ENTRENAMIENTO PRINCIPAL ==========
results = []
tickers = df['ticker'].unique()
for ticker in tickers:
    df_t = df[df['ticker'] == ticker].copy().sort_values('fecha')
    df_t.set_index('fecha', inplace=True)
    print(f'\n================  {ticker}  ================')
    for target in TARGET_VARS:
        series = df_t[target].dropna()
        if len(series) < 40:
            print(f'⚠️  {ticker} - {target}: muy pocos datos, omitiendo.')
            continue
        print(f'→ Entrenando modelo para {ticker} - {target} (n={len(series)})')
        try:
            best, forecast_df = train_forecast(series, fh=FH, modelos_a_probar=MODELOS_A_PROBAR)
            best_name = str(best) if not hasattr(best, 'model') else type(best.model).__name__
            for f_date, row in forecast_df.iterrows():
                results.append({
                    'ticker': ticker,
                    'variable': target,
                    'date': f_date,
                    'forecast': row['y_pred'],
                    'lower': row.get('y_pred_lower', pd.NA),
                    'upper': row.get('y_pred_upper', pd.NA)
                })
            plot_forecast(series, forecast_df, f'{ticker} - {target}', zoom_window=30, model_name=best_name)
        except Exception as e:
            logger.error(f'Error entrenando {ticker}-{target}: {e}')
            continue

# ========== GUARDADO DE RESULTADOS ==========
pred_df = pd.DataFrame(results)
pred_df.to_csv(OUT_DIR/'forecast_full.csv', index=False)
logger.info(f'Se guardó forecast_full.csv con {len(pred_df)} filas')

# ========== REPORTE DOCX ==========
try:
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    doc.add_heading('Pronóstico de Acciones', level=1)
    doc.add_paragraph(f'Ejecución: {RUN_TS}')
    for ticker in pred_df['ticker'].unique():
        doc.add_heading(ticker, level=2)
        sub = pred_df[pred_df['ticker']==ticker].copy()
        for var in sub['variable'].unique():
            sub2 = sub[sub['variable']==var]
            doc.add_heading(var, level=3)
            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Fecha'
            hdr_cells[1].text = 'Predicción'
            hdr_cells[2].text = 'Lower'
            hdr_cells[3].text = 'Upper'
            for _,r in sub2.iterrows():
                # Soporte para Period, Timestamp, str, etc.
                fecha = r['date']
                if hasattr(fecha, 'strftime'):
                    fecha_str = fecha.strftime('%Y-%m-%d')
                else:
                    fecha_str = str(fecha)
                row_cells = table.add_row().cells
                row_cells[0].text = fecha_str
                row_cells[1].text = f"{r['forecast']:.2f}" if pd.notna(r['forecast']) else '-'
                row_cells[2].text = '-' if pd.isna(r['lower']) else f"{r['lower']:.2f}"
                row_cells[3].text = '-' if pd.isna(r['upper']) else f"{r['upper']:.2f}"
            # Agregar imágenes si existen
            safe_title = f"{ticker} - {var}".replace(" / ", "_")
            img_full = OUT_DIR / f"{safe_title}_full.png"
            img_zoom = OUT_DIR / f"{safe_title}_zoom.png"
            if img_full.exists():
                doc.add_paragraph('Gráfico completo:')
                doc.add_picture(str(img_full), width=Inches(5.5))
            if img_zoom.exists():
                doc.add_paragraph(f'Gráfico últimos 30 días + forecast:')
                doc.add_picture(str(img_zoom), width=Inches(5.5))
    doc.save(OUT_DIR/'reporte_proyecciones.docx')
    logger.info('Se generó reporte_proyecciones.docx')
except ImportError:
    logger.warning('python-docx no instalado; omitiendo reporte DOCX')
except Exception as e:
    logger.error(f'Error generando DOCX: {e}')

print('Proceso finalizado. Revisa el directorio:', OUT_DIR)
